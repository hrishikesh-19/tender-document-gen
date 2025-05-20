import streamlit as st
from google import genai
from google.genai import types
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import uuid
import os
import re
import datetime

# Optional: PDF support (Tab 2)
import fitz  # PyMuPDF

# Initialize Gemini client
client = genai.Client(api_key=st.secrets["gemini_api_key"])

# Load system instruction for tender drafting
with open("tender_prompt.txt", "r") as file:
    sys_instruct = file.read()

st.title("📄 AI Tender Document Generator")

# Tab layout
tab1, tab2 = st.tabs(["📄 Generate From Scratch", "📝 Modify Existing Tender"])

# --------------------- TAB 1 ---------------------
with tab1:
    st.sidebar.header("Tender Metadata")
    tender_title = st.sidebar.text_input("Tender Title", "AI-Based Digital Infrastructure")
    tender_number = st.sidebar.text_input("Tender Number", "TDR-2024-001")
    issue_date = st.sidebar.date_input("Issue Date", datetime.date.today())

    def get_prompt_suggestions(user_input, ai_response):
        prompt = f"""
You are a helpful AI assistant that helps users draft professional tender documents. Based on the user's latest input and the AI's response, suggest 3 to 5 logical follow-up prompts or sections the user might want to include next.

Format your output as a Python list of strings.

User Input:
\"\"\"{user_input}\"\"\"

AI Response:
\"\"\"{ai_response}\"\"\"

Return only the list like:
["Add payment terms", "Include evaluation criteria", "Mention timeline and deadlines", "List eligibility criteria"]
"""
        suggestion_chat = client.chats.create(
            model="gemini-2.0-flash",
            config=types.GenerateContentConfig(
                system_instruction="You are a prompt suggestion expert for AI-generated tender documents.",
                response_mime_type="application/json"
            )
        )
        result = suggestion_chat.send_message(prompt)
        try:
            suggestions = eval(result.text.strip())
            if isinstance(suggestions, list):
                return suggestions
        except Exception as e:
            print(f"Error parsing suggestions: {e}")
        return ["Include scope of work", "Define bidder qualifications", "Mention deliverables"]

    def extract_placeholder_values_from_input(placeholders, user_input):
        prompt = f"""
You are a helpful assistant that extracts structured placeholder values from user messages.

Given this list of placeholders:
{placeholders}

And this user message:
\"\"\"{user_input}\"\"\"

Return a JSON object mapping placeholder names to their values, like:
{{
  "Deadline": "31 May 2025",
  "Bid Amount": "50000 INR"
}}

Only return JSON. No explanation.
"""
        try:
            chat = client.chats.create(
                model="gemini-2.0-flash",
                config=types.GenerateContentConfig(
                    system_instruction="You convert natural language inputs into structured placeholder-value pairs.",
                    response_mime_type="application/json"
                )
            )
            result = chat.send_message(prompt)
            return eval(result.text.strip())
        except Exception as e:
            print(f"AI mapping error: {e}")
            return {}

    def extract_placeholders(text):
        matches = re.findall(r'\[.*?\]|\{.*?\}|\<.*?\>', text)
        return sorted(set([m.strip("[]{}<>") for m in matches]))

    def add_page_number(section):
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run = paragraph.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def generate_formatted_tender_doc(messages):
        doc = Document()
        doc.add_paragraph()
        title = doc.add_heading("Tender Document", 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(f"Tender Title: {tender_title}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Tender Number: {tender_number}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f"Issue Date: {issue_date.strftime('%d-%m-%Y')}").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_page_break()

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        def clean_markdown(text):
                text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # bold
                text = re.sub(r"\*(.*?)\*", r"\1", text)      # italic
                text = re.sub(r"`(.*?)`", r"\1", text)        # inline code
                text = re.sub(r"#+\s", "", text)              # remove markdown headings like ## Header
                return text.strip()
        
        for msg in st.session_state.messages:
            if msg["role"] == "assistant":
                content = msg["content"].strip()
                for line in content.split("\n"):
                    line = clean_markdown(line.strip())
                    if not line:
                        continue
                    if line.endswith(":") and len(line.split()) < 10:
                        doc.add_heading(line.rstrip(":"), level=1)
                    elif re.match(r"^[-\u2022]\s", line):
                        doc.add_paragraph(line[2:].strip(), style='List Bullet')
                    else:
                        doc.add_paragraph(line)

        section = doc.sections[0]
        footer_para = section.footer.paragraphs[0]
        footer_para.text = "Confidential - Generated via AI Tender Assistant"
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        add_page_number(section)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # --- Session state ---
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())

    if "chat_session" not in st.session_state:
        st.session_state.chat_session = client.chats.create(
            model="gemini-2.0-flash",
            config=types.GenerateContentConfig(system_instruction=sys_instruct),
        )

    if "messages" not in st.session_state:
        st.session_state.messages = []

    if "selected_prompt" not in st.session_state:
        st.session_state.selected_prompt = None

    if "last_response" not in st.session_state:
        st.session_state.last_response = ""

    if "suggestions" not in st.session_state:
        st.session_state.suggestions = []

    # --- Initial greeting ---
    if not st.session_state.messages:
        st.session_state.messages.append({
            "role": "assistant",
            "content": "Hello! I can help you draft a professional tender document. Tell me your requirement to get started."
        })

    # --- Chat History ---
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    # --- Chat Input ---
    user_input = st.chat_input("Describe what the tender is for or fill in the placeholders...")

    if user_input:
        with st.chat_message("user"):
            st.markdown(user_input)
        st.session_state.messages.append({"role": "user", "content": user_input})

        with st.spinner("Drafting section..."):
            response = st.session_state.chat_session.send_message(user_input)
            bot_response = response.text

        prev_placeholders = extract_placeholders(st.session_state.last_response)
        if prev_placeholders:
            extracted_values = extract_placeholder_values_from_input(prev_placeholders, user_input)
            if extracted_values:
                updated = st.session_state.last_response
                for ph, val in extracted_values.items():
                    pattern = rf"[\[\{{\<]{ph}[\]\}}\>]"
                    updated = re.sub(pattern, str(val), updated, flags=re.IGNORECASE)

                st.session_state.messages[-1]["content"] = updated
                st.session_state.last_response = updated

                with st.chat_message("assistant"):
                    st.success("✅ Updated section with your inputs:")
                    st.markdown(updated)
                st.rerun()

        st.session_state.last_response = bot_response
        st.session_state.messages.append({"role": "assistant", "content": bot_response})
        with st.chat_message("assistant"):
            st.markdown(bot_response)

        st.session_state.suggestions = get_prompt_suggestions(user_input, bot_response)
        st.rerun()

    if st.session_state.selected_prompt:
        prompt = st.session_state.selected_prompt
        with st.chat_message("user"):
            st.markdown(prompt)
        st.session_state.messages.append({"role": "user", "content": prompt})

        with st.spinner("Generating content..."):
            response = st.session_state.chat_session.send_message(prompt)
            bot_response = response.text

        st.session_state.last_response = bot_response
        st.session_state.messages.append({"role": "assistant", "content": bot_response})
        with st.chat_message("assistant"):
            st.markdown(bot_response)

        st.session_state.suggestions = get_prompt_suggestions(prompt, bot_response)
        st.session_state.selected_prompt = None
        st.rerun()

    if st.session_state.suggestions and not st.session_state.selected_prompt:
        with st.chat_message("assistant"):
            st.markdown("Would you like to include:")
            for i, sug in enumerate(st.session_state.suggestions):
                if st.button(sug, key=f"suggest_{i}_{st.session_state.session_id}"):
                    st.session_state.selected_prompt = sug
                    st.rerun()

    if st.session_state.last_response:
        placeholders = extract_placeholders(st.session_state.last_response)
        if placeholders:
            with st.chat_message("assistant"):
                st.info(f"This section includes placeholders: {', '.join(placeholders)}. You can mention the values in chat and I’ll replace them automatically.")

    if any(m["role"] == "assistant" for m in st.session_state.messages):
        word_file = generate_formatted_tender_doc(st.session_state.messages)
        st.download_button(
            label="📥 Download Full Tender Document (.docx)",
            data=word_file,
            file_name="AI_Generated_Tender_Document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --------------------- TAB 2 ---------------------
with tab2:
    st.subheader("📝 Upload and Chat to Modify Existing Tender Document")

    uploaded_file = st.file_uploader("Upload a Tender Document (.docx or .pdf)", type=["docx", "pdf"])
    extracted_text = ""

    if uploaded_file and "tab2_uploaded" not in st.session_state:
        if uploaded_file.name.endswith(".pdf"):
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            extracted_text = "\n".join([page.get_text() for page in doc])
        else:
            doc = Document(uploaded_file)
            extracted_text = "\n".join([para.text for para in doc.paragraphs])

        st.session_state.tab2_uploaded = extracted_text

        # Step 1: AI Suggestions
        with st.spinner("Analyzing document and suggesting edits..."):
            suggestion_prompt = f"""
You're an AI that reviews tender documents and suggests improvements.

Here is the tender content:
\"\"\"{extracted_text}\"\"\"

Suggest 3 to 5 specific improvements in this format:
- Suggestion 1
- Suggestion 2
"""
            suggestion_chat = client.chats.create(
                model="gemini-2.0-flash",
                config=types.GenerateContentConfig(
                    system_instruction="You review tenders and suggest improvements.",
                    response_mime_type="text/plain"
                )
            )
            suggestion_response = suggestion_chat.send_message(suggestion_prompt)
            suggestions = suggestion_response.text.strip()
            st.session_state.tab2_suggestions = [s.lstrip("- ").strip() for s in suggestions.split("\n") if s.strip()]

        st.success("✅ Suggestions loaded. Start chatting to apply changes.")

        # Step 2: Initialize chat session & messages
        st.session_state.tab2_chat = client.chats.create(
            model="gemini-2.0-flash",
            config=types.GenerateContentConfig(
                system_instruction=f"You are editing a tender document based on user instructions. The current content is:\n\n'''{st.session_state.tab2_uploaded}'''\n\nApply any user changes and return the full updated document.",
                response_mime_type="text/plain"
            )
        )
        st.session_state.tab2_messages = []

    # --- If document uploaded and chat ready ---
    if "tab2_chat" in st.session_state:
        st.info("💡 Gemini Suggestions:")
        for suggestion in st.session_state.tab2_suggestions:
            if st.button(suggestion, key=f"suggest_{suggestion}"):
                # Act as if user entered this message in chat
                with st.chat_message("user"):
                    st.markdown(suggestion)
                st.session_state.tab2_messages.append({"role": "user", "content": suggestion})

                with st.spinner("Applying suggestion..."):
                    response = st.session_state.tab2_chat.send_message(suggestion)
                    ai_reply = response.text.strip()

                with st.chat_message("assistant"):
                    st.markdown(ai_reply)
                st.session_state.tab2_messages.append({"role": "assistant", "content": ai_reply})
                st.session_state.tab2_last_updated_doc = ai_reply
                st.rerun()

        # Show all previous chat messages
        for msg in st.session_state.tab2_messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        user_input = st.chat_input("Describe your change...")
        if user_input:
            with st.chat_message("user"):
                st.markdown(user_input)
            st.session_state.tab2_messages.append({"role": "user", "content": user_input})

            with st.spinner("Applying changes..."):
                response = st.session_state.tab2_chat.send_message(user_input)
                ai_reply = response.text.strip()

            with st.chat_message("assistant"):
                st.markdown(ai_reply)
            st.session_state.tab2_messages.append({"role": "assistant", "content": ai_reply})
            st.session_state.tab2_last_updated_doc = ai_reply
            st.rerun()

        # Show export button if we have a final draft
        if "tab2_last_updated_doc" in st.session_state:
            def clean_markdown(text):
                text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)  # bold
                text = re.sub(r"\*(.*?)\*", r"\1", text)      # italic
                text = re.sub(r"`(.*?)`", r"\1", text)        # inline code
                text = re.sub(r"#+\s", "", text)              # remove markdown headings like ## Header
                return text.strip()
            
            def generate_docx_from_text(text):
                doc = Document()
                for para in text.split("\n"):
                    para = clean_markdown(para.strip())
                    doc.add_paragraph(para)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                return buffer


            final_docx = generate_docx_from_text(st.session_state.tab2_last_updated_doc)

            st.download_button(
                label="📥 Download Final Updated Document (.docx)",
                data=final_docx,
                file_name="Final_Updated_Tender.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

